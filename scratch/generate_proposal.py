import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_proposal():
    doc = docx.Document()
    
    # Set Margin (Standard Indonesian Academic: Top 3, Bottom 3, Left 4, Right 3 cm)
    # 1 Inch = 2.54 cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.18)      # ~3cm
        section.bottom_margin = Inches(1.18)   # ~3cm
        section.left_margin = Inches(1.57)     # ~4cm
        section.right_margin = Inches(1.18)    # ~3cm
        
    # Styles config
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)
    
    # Helper to add headings
    def add_chapter_heading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)
        
    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    # ==================== COVER PAGE ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(48)
    
    title_run = title_p.add_run(
        "DRAFT PROPOSAL TUGAS AKHIR\n\n"
        "RANCANG BANGUN SISTEM MONITORING IOT MANUFAKTUR BERBASIS RASPBERRY PI\n"
        "DENGAN FITUR WEB CONTROL PANEL SERVER DAN PENGAMANAN\n"
        "ENKRIPSI SSL JARINGAN LOKAL"
    )
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Times New Roman'
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Sub cover info
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_before = Pt(72)
    info_p.paragraph_format.space_after = Pt(120)
    info_run = info_p.add_run(
        "Diajukan sebagai Rancangan Penelitian Tugas Akhir\n"
        "Program Studi Teknik Komputer / Teknik Informatika"
    )
    info_run.font.size = Pt(12)
    info_run.font.name = 'Times New Roman'
    info_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Author info placeholder
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run(
        "Disusun Oleh:\n"
        "Nama Mahasiswa / NIM\n\n\n"
        "FAKULTAS TEKNIK / JURUSAN TEKNIK INFORMATIKA\n"
        "TAHUN 2026"
    )
    author_run.bold = True
    author_run.font.size = Pt(12)
    author_run.font.name = 'Times New Roman'
    author_run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_page_break()

    # ==================== BAB I ====================
    add_chapter_heading("BAB I\nPENDAHULUAN")
    
    add_section_heading("1.1 Latar Belakang Masalah")
    doc.add_paragraph(
        "Lini produksi dan perakitan pada industri manufaktur dituntut untuk senantiasa memiliki efisiensi tinggi serta tingkat "
        "kesalahan pengemasan yang mendekati nol persen (zero defect). Salah satu stasiun vital dalam alur manufaktur ini adalah "
        "stasiun pemeriksaan kualitas (Quality Control/QC) dan pengemasan produk. Verifikasi kelengkapan komponen di dalam "
        "kemasan sebelum didistribusikan umumnya memanfaatkan kombinasi sensor berat (load cell) dan pengolahan citra kamera (computer vision) "
        "berbasis Internet of Things (IoT). Pada arsitektur lokal pabrik, komputer papan tunggal (single-board computer) seperti Raspberry Pi "
        "kerap digunakan sebagai server mini lokal karena efisiensi energi, ukuran yang kompak, serta kemampuannya menangani server web mandiri."
    )
    
    doc.add_paragraph(
        "Penelitian Tugas Akhir ini merupakan kelanjutan pengembangan dari proyek Project-Based Learning (PBL) yang telah dikerjakan oleh "
        "peneliti sebelumnya pada Semester 4. Pada proyek PBL awal tersebut, sistem yang dibangun berfokus pada rancang bangun "
        "sistem pelacakan produk manufaktur sederhana menggunakan kamera webcam untuk mendeteksi kelengkapan isi box berdasarkan "
        "pengolahan citra lokal, dengan penyimpanan data pada database lokal yang terhubung melalui protokol HTTP standar."
    )
    
    doc.add_paragraph(
        "Meskipun sistem pelacakan awal hasil PBL tersebut berhasil dijalankan, terdapat beberapa batasan sistem yang memerlukan "
        "pengembangan lebih lanjut agar memenuhi standar industri yang aman dan andal. Batasan pertama adalah tidak adanya integrasi "
        "verifikasi fisik tambahan seperti sensor berat untuk menjamin presisi kualitas isi box. Batasan kedua adalah aspek keamanan "
        "transaksi data sensor dan login operator yang masih berjalan tanpa enkripsi (HTTP biasa), sehingga rentan terhadap serangan "
        "penyadapan data (sniffing) di jaringan lokal. Batasan ketiga adalah pengelolaan fisik server Raspberry Pi yang sepenuhnya masih "
        "bergantung pada baris perintah CLI Linux melalui protokol SSH, menyulitkan operator pabrik non-IT jika terjadi kendala teknis."
    )
    
    doc.add_paragraph(
        "Untuk mengatasi batasan-batasan dari proyek PBL tersebut, penelitian Tugas Akhir ini merancang stasiun verifikasi ganda "
        "dengan menambahkan sensor berat (Load Cell HX711) yang terintegrasi. Selain itu, sistem dikembangkan ke arah infrastruktur "
        "jaringan industri yang aman dengan menerapkan enkripsi SSL/TLS lokal (HTTPS) guna mencegah sniffing, pembatasan akses remote via VPN, "
        "serta pembuatan modul Web Server Control Panel (WSCP) berbasis Laravel pada dashboard admin. Dengan modul ini, operator dapat dengan "
        "mudah memantau kesehatan hardware Raspberry Pi (CPU, suhu, RAM) dan mengontrol layanan server secara visual langsung dari web."
    )
    
    add_section_heading("1.2 Perumusan Masalah")
    doc.add_paragraph(
        "Rumusan masalah yang dibahas pada penelitian Tugas Akhir ini adalah:\n"
        "1. Bagaimana merancang arsitektur jaringan lokal terenkripsi HTTPS (SSL/TLS) lokal pada Raspberry Pi untuk mengamankan data transaksi sensor IoT dari ancaman penyadapan siber (sniffing)?\n"
        "2. Bagaimana membangun modul Web Server Control Panel (WSCP) berbasis Laravel pada dashboard admin guna mempermudah pemantauan kinerja perangkat keras (hardware resource) dan pengelolaan layanan server Raspberry Pi?\n"
        "3. Bagaimana mengintegrasikan stasiun verifikasi kamera (blob detection) dan sensor berat (HX711) ke dalam sistem monitoring terenkripsi tersebut?"
    )
    
    add_section_heading("1.3 Batasan Masalah")
    doc.add_paragraph(
        "Penelitian Tugas Akhir ini memiliki batasan-batasan masalah sebagai berikut:\n"
        "1. Keamanan siber yang dikembangkan difokuskan pada enkripsi data transaksi lokal menggunakan protokol HTTPS SSL/TLS lokal dan autentikasi jaringan VPN privat.\n"
        "2. Perangkat keras server lokal dibatasi pada komputer papan tunggal Raspberry Pi 4 dengan sistem operasi berbasis Debian Linux.\n"
        "3. Pengembangan aplikasi web server dan panel manajemen kontrol server (WSCP) dibangun menggunakan framework Laravel 11.\n"
        "4. Sensor verifikasi fisik dibatasi pada kamera webcam USB lokal dan timbangan digital meja menggunakan load cell yang terhubung ke modul penguat HX711."
    )
    
    add_section_heading("1.4 Tujuan Penelitian")
    doc.add_paragraph(
        "Tujuan yang ingin dicapai melalui penelitian Tugas Akhir ini adalah:\n"
        "1. Meningkatkan keamanan pengiriman data hasil sensor dan kredensial login di area pabrik dengan mengimplementasikan sertifikat enkripsi SSL lokal pada Nginx.\n"
        "2. Menyediakan modul Web Server Control Panel (WSCP) interaktif berbasis web untuk memudahkan pemeliharaan server lokal tanpa memerlukan CLI.\n"
        "3. Membangun purwarupa sistem verifikasi ganda (QC visual kamera & berat barang) yang datanya terkirim secara aman ke database pusat secara real-time."
    )
    
    add_section_heading("1.5 Manfaat Penelitian")
    doc.add_paragraph(
        "Penelitian ini diharapkan memberikan manfaat sebagai berikut:\n"
        "1. Manfaat bagi Industri Manufaktur: Menyediakan sistem verifikasi kelengkapan box yang aman dari manipulasi data pihak luar serta memangkas waktu pemeliharaan server IoT lokal.\n"
        "2. Manfaat bagi Peneliti dan Akademisi: Menjadi kajian pustaka mengenai implementasi integrasi antara IoT edge computing, web development Laravel, dan pengamanan administrasi jaringan siber lokal."
    )
    
    doc.add_page_break()

    # ==================== BAB II ====================
    add_chapter_heading("BAB II\nTINJAUAN PUSTAKA")
    
    add_section_heading("2.1 Internet of Things (IoT) di Sektor Manufaktur")
    doc.add_paragraph(
        "Internet of Things (IoT) merujuk pada jaringan objek fisik yang ditanami sensor, perangkat lunak, dan teknologi komunikasi "
        "untuk saling bertukar data dengan perangkat lain. Pada manufaktur modern (Industri 4.0), IoT diimplementasikan untuk pelacakan "
        "produk secara real-time di lini perakitan guna memastikan kualitas kemasan dan meminimalkan kesalahan kemasan secara instan."
    )
    
    add_section_heading("2.2 Keamanan Jaringan Enkripsi SSL/TLS Lokal")
    doc.add_paragraph(
        "Secure Sockets Layer (SSL) dan Transport Layer Security (TLS) adalah protokol kriptografi yang dirancang untuk memberikan "
        "keamanan komunikasi melalui jaringan komputer. Dengan memasang sertifikat SSL lokal (Self-Signed) pada server Nginx di Raspberry Pi, "
        "koneksi HTTP standar pada port 8000 dan 8001 ditingkatkan menjadi HTTPS. Seluruh data transaksi yang dikirimkan oleh browser "
        "klien, termasuk data sensor dan kata sandi login, akan diacak menggunakan kunci enkripsi simetris sehingga kebal dari aksi sadap data (sniffing)."
    )
    
    add_section_heading("2.3 Web Server Control Panel (WSCP)")
    doc.add_paragraph(
        "Web Server Control Panel (WSCP) adalah antarmuka grafis berbasis peramban web yang menyederhanakan tugas administrasi server "
        "seperti pemantauan beban perangkat keras, modifikasi file konfigurasi, dan kontrol sistem (systemd service). Pemanfaatan bahasa "
        "pemrograman PHP dengan fungsi eksekusi perintah shell (seperti shell_exec) memungkinkan aplikasi web memanggil utilitas sistem operasi "
        "Linux secara aman untuk membaca status server atau melakukan tindakan pemulihan sistem."
    )
    
    doc.add_page_break()

    # ==================== BAB III ====================
    add_chapter_heading("BAB III\nMETODOLOGI PENELITIAN")
    
    add_section_heading("3.1 Metode Pengembangan Sistem")
    doc.add_paragraph(
        "Metodologi pengembangan sistem yang digunakan dalam penelitian ini adalah metode Prototyping untuk rekayasa perangkat lunak "
        "aplikasi dan metode NDLC (Network Development Life Cycle) untuk perancangan dan implementasi infrastruktur jaringan aman."
    )
    
    add_section_heading("3.2 Perancangan Topologi dan Arsitektur Jaringan")
    doc.add_paragraph(
        "Raspberry Pi dikonfigurasi sebagai Gateway Router dan Access Point mandiri (SSID: ProTrack-WiFi) dengan subnet IP 10.42.0.1. "
        "Nginx bertindak sebagai web server sekaligus Reverse Proxy untuk meneruskan lalu lintas HTTPS port 443 secara internal ke port aplikasi "
        "Laravel 8000 (Operator) dan 8001 (Admin). Jalur VPN WireGuard dikonfigurasi di sisi Raspberry Pi untuk menerima koneksi "
        "enkripsi tunneling dari perangkat eksternal admin guna melakukan pemeliharaan server jarak jauh secara privat."
    )
    
    add_section_heading("3.3 Rencana Pengujian dan Parameter Keberhasilan")
    doc.add_paragraph(
        "Pengujian kualitas sistem akan dikelompokkan ke dalam tiga parameter uji utama:\n"
        "1. Pengujian Fungsionalitas QC: Mengukur akurasi deteksi webcam (blob detection) dan sensor timbangan berat load cell HX711.\n"
        "2. Pengujian Kinerja Server (WSCP): Mengukur respons waktu eksekusi kontrol sistem (systemd restart Nginx/Laravel) dari antarmuka web.\n"
        "3. Pengujian Keamanan Jaringan: Melakukan pembuktian penetrasi penyadapan jaringan nirkabel menggunakan Wireshark untuk menganalisis data payload paket HTTP vs HTTPS."
    )

    # Save document
    doc.save("c:/Users/ikhwa/operator-system/draft_proposal_ta.docx")
    print("Document successfully created!")

if __name__ == "__main__":
    create_proposal()
